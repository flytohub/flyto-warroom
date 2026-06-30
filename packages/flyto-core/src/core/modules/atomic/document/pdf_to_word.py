# Copyright 2026 Flyto2. Licensed under Apache-2.0. See LICENSE.

"""
PDF to Word Converter Module
Convert PDF files to Word documents (.docx)
"""
import logging
import os
from typing import Any, Dict

from ...registry import register_module
from ...schema import compose, presets


logger = logging.getLogger(__name__)


@register_module(
    module_id='pdf.to_word',
    version='1.0.0',
    category='document',
    subcategory='pdf',
    tags=['pdf', 'word', 'docx', 'convert', 'document', 'path_restricted'],
    label='PDF to Word',
    label_key='modules.pdf.to_word.label',
    description='Convert PDF files to Word documents (.docx)',
    description_key='modules.pdf.to_word.description',
    icon='FileOutput',
    color='#2563EB',

    # Connection types
    input_types=['file_path'],
    output_types=['file_path'],
    can_connect_to=['file.*', 'pdf.*', 'flow.*'],
    can_receive_from=['file.*', 'data.*', 'http.*', 'flow.*', 'start'],

    # Execution settings
    timeout_ms=300000,
    retryable=False,
    concurrent_safe=True,

    # Security settings
    requires_credentials=False,
    handles_sensitive_data=True,
    required_permissions=['filesystem.read', 'filesystem.write'],

    params_schema=compose(
        presets.DOC_INPUT_PATH(placeholder='/path/to/document.pdf'),
        presets.DOC_OUTPUT_PATH(placeholder='/path/to/output.docx'),
        presets.DOC_PRESERVE_FORMATTING(),
        presets.DOC_PAGES(),
    ),
    output_schema={
        'output_path': {
            'type': 'string',
            'description': 'Path to the generated Word document'
        ,
                'description_key': 'modules.pdf.to_word.output.output_path.description'},
        'page_count': {
            'type': 'number',
            'description': 'Number of pages converted'
        ,
                'description_key': 'modules.pdf.to_word.output.page_count.description'},
        'file_size': {
            'type': 'number',
            'description': 'Size of the output file in bytes'
        ,
                'description_key': 'modules.pdf.to_word.output.file_size.description'}
    },
    examples=[
        {
            'title': 'Convert entire PDF to Word',
            'title_key': 'modules.pdf.to_word.examples.basic.title',
            'params': {
                'input_path': '/tmp/document.pdf'
            }
        },
        {
            'title': 'Convert specific pages',
            'title_key': 'modules.pdf.to_word.examples.pages.title',
            'params': {
                'input_path': '/tmp/document.pdf',
                'output_path': '/tmp/output.docx',
                'pages': '1-5'
            }
        }
    ],
    author='Flyto Team',
    license='MIT'
)
async def pdf_to_word(context: Dict[str, Any]) -> Dict[str, Any]:
    """Convert PDF to Word document"""
    try:
        import pypdf
    except ImportError:
        raise ImportError("pypdf is required. Install with: pip install pypdf")
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    params = context['params']
    input_path = params['input_path']
    pages_param = params.get('pages', 'all')
    preserve_formatting = params.get('preserve_formatting', True)
    output_path = _resolve_pdf_output(params, input_path)

    if not os.path.exists(input_path):
        raise FileNotFoundError(f"PDF file not found: {input_path}")
    _ensure_output_dir(output_path)

    if ".." in input_path:
        raise Exception("Invalid file path")
    with open(input_path, 'rb') as f:
        reader = pypdf.PdfReader(f)
        total_pages = len(reader.pages)
        page_indices = _parse_page_range(pages_param, total_pages)
        doc = Document()
        _add_pdf_title(doc, reader, WD_PARAGRAPH_ALIGNMENT)
        converted_pages = _convert_pages(
            doc, reader, page_indices, total_pages, preserve_formatting, Pt
        )

    doc.save(output_path)
    file_size = os.path.getsize(output_path)
    logger.info(f"Converted PDF to Word: {input_path} -> {output_path} ({converted_pages} pages)")

    return {
        'ok': True,
        'output_path': output_path,
        'page_count': converted_pages,
        'total_pages': total_pages,
        'file_size': file_size,
        'message': f'Successfully converted {converted_pages} pages to Word document'
    }


def _resolve_pdf_output(params: Dict[str, Any], input_path: str) -> str:
    output_path = params.get('output_path')
    if not output_path:
        base = os.path.splitext(input_path)[0]
        output_path = f"{base}.docx"
    return output_path


def _ensure_output_dir(output_path: str):
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)


def _add_pdf_title(doc, reader, alignment_enum):
    if reader.metadata and reader.metadata.get('/Title'):
        title = doc.add_heading(reader.metadata.get('/Title'), level=0)
        title.alignment = alignment_enum.CENTER


def _convert_pages(doc, reader, page_indices, total_pages, preserve_formatting, Pt):
    converted_pages = 0
    for idx in page_indices:
        if 0 <= idx < total_pages:
            text = (reader.pages[idx].extract_text() or "").strip()
            if not text:
                continue

            if len(page_indices) > 1:
                doc.add_heading(f"Page {idx + 1}", level=2)

            for para_text in text.split('\n\n'):
                para_text = para_text.strip()
                if para_text:
                    para = doc.add_paragraph(para_text)
                    if preserve_formatting:
                        for run in para.runs:
                            run.font.size = Pt(11)

            converted_pages += 1

            if idx != page_indices[-1]:
                doc.add_page_break()

    return converted_pages


def _parse_page_range(pages: str, total: int) -> list:
    """Parse page range string to list of indices (0-based)"""
    if pages.lower() == 'all':
        return list(range(total))

    indices = []
    parts = pages.replace(' ', '').split(',')

    for part in parts:
        if '-' in part:
            start, end = part.split('-')
            start_idx = int(start) - 1
            end_idx = int(end)
            indices.extend(range(start_idx, min(end_idx, total)))
        else:
            idx = int(part) - 1
            if 0 <= idx < total:
                indices.append(idx)

    return sorted(set(indices))
